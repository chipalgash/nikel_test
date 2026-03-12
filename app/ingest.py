from __future__ import annotations

import re
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document

from .config import SETTINGS
from .schemas import Chunk

logger = logging.getLogger(__name__)


@dataclass
class Block:
    text: str
    source_path: str
    source_file: str
    section_path: str
    block_type: str


SECTION_PATTERNS = [
    re.compile(r"^раздел\s+[ivxlcdm0-9]+", re.IGNORECASE),
    re.compile(r"^глава\s+[0-9.]+", re.IGNORECASE),
    re.compile(r"^статья\s+[0-9.\-]+", re.IGNORECASE),
    re.compile(r"^пункт\s+[0-9.\-]+", re.IGNORECASE),
    re.compile(r"^[0-9]+(\.[0-9]+){0,3}\s+"),
    re.compile(r"^приложение\s+[а-яa-z0-9]+", re.IGNORECASE),
    re.compile(r"^таблица\s+[0-9.\-]+", re.IGNORECASE),
]


def is_heading(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if len(t) <= 4:
        return False
    return any(p.search(t) for p in SECTION_PATTERNS)


def _keep_heading_as_content(text: str) -> bool:
    """Keep long numbered clauses (e.g. 6.1.10 ...) as searchable content."""
    t = normalize_text(text)
    if not t:
        return False
    # Normative clauses often start with numbering and contain a full sentence.
    if re.match(r"^[0-9]+(\.[0-9]+){1,3}\s+", t):
        low = t.lower()
        if len(t) >= 48 and any(
            x in low for x in ("следует", "должен", "должна", "принимать", "принима", "приведен", "допустим")
        ):
            return True
        if len(t) >= 56 and re.search(r"(мм|км/ч|°c|таблиц|пункт|отклонени)", low):
            return True
    return False


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _normalize_table_key(text: str, idx: int) -> str:
    t = normalize_text(text).lower()
    t = re.sub(r"[^a-zа-я0-9]+", "_", t, flags=re.IGNORECASE).strip("_")
    return t or f"col_{idx}"


def _has_table_header(cells: List[str]) -> bool:
    if not cells:
        return False
    # Header rows are usually text-heavy and not dominated by numbers.
    non_empty = [c for c in cells if c]
    if not non_empty:
        return False
    digit_ratio = sum(1 for c in non_empty if re.search(r"\d", c)) / len(non_empty)
    return digit_ratio < 0.6


def _table_row_text(
    table_id: int,
    row_id: int,
    headers: List[str],
    cells: List[str],
) -> str:
    raw_cells = [c for c in cells if c]
    row_raw = " | ".join(raw_cells)
    kv_parts: List[str] = []
    for i, cell in enumerate(cells):
        cell = normalize_text(cell)
        if not cell:
            continue
        key_src = headers[i] if i < len(headers) else f"col_{i+1}"
        key = _normalize_table_key(key_src, i + 1)
        kv_parts.append(f"{key}={cell}")

    nums = re.findall(r"\d+(?:[.,]\d+)?", row_raw)
    nums_text = ",".join(nums) if nums else ""
    header_text = " | ".join([normalize_text(h) for h in headers if normalize_text(h)])
    return (
        f"[TABLE_ROW] [TABLE_ID={table_id}] [ROW_ID={row_id}] "
        f"[HEADER] {header_text} "
        f"[ROW_RAW] {row_raw} "
        f"[ROW_KV] {' ; '.join(kv_parts)} "
        f"[NUMS] {nums_text}"
    ).strip()


def parse_docx(docx_path: Path) -> List[Block]:
    logger.info("Parsing DOCX: %s", docx_path.name)
    doc = Document(str(docx_path))
    source_path = str(docx_path.resolve())
    source_file = docx_path.name

    section_stack: List[str] = []
    blocks: List[Block] = []

    for p in doc.paragraphs:
        raw = normalize_text(p.text)
        if not raw:
            continue

        if is_heading(raw):
            section_stack.append(raw)
            if len(section_stack) > 4:
                section_stack = section_stack[-4:]
            if _keep_heading_as_content(raw):
                blocks.append(
                    Block(
                        text=raw,
                        source_path=source_path,
                        source_file=source_file,
                        section_path=" > ".join(section_stack),
                        block_type="paragraph",
                    )
                )
            continue

        blocks.append(
            Block(
                text=raw,
                source_path=source_path,
                source_file=source_file,
                section_path=" > ".join(section_stack),
                block_type="paragraph",
            )
        )

    for t_idx, table in enumerate(doc.tables, start=1):
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        headers = rows[0] if _has_table_header(rows[0]) else [f"col_{i+1}" for i in range(len(rows[0]))]
        for r_idx, cells in enumerate(rows, start=1):
            if not any(cells):
                continue
            row_text = _table_row_text(table_id=t_idx, row_id=r_idx, headers=headers, cells=cells)
            blocks.append(
                Block(
                    text=row_text,
                    source_path=source_path,
                    source_file=source_file,
                    section_path=" > ".join(section_stack),
                    block_type="table_row",
                )
            )

    logger.info("Parsed DOCX: %s | blocks=%s", docx_path.name, len(blocks))
    return blocks


def _split_child_text(text: str, child_size: int, child_overlap: int) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= child_size:
        return [text]

    parts: List[str] = []
    start = 0
    step = max(80, child_size - child_overlap)
    while start < len(text):
        end = min(len(text), start + child_size)
        piece = text[start:end].strip()
        if piece:
            parts.append(piece)
        if end >= len(text):
            break
        start += step
    return parts


def _build_table_context_chunks(
    blocks: List[Block],
    source_path: str,
    source_file: str,
    start_chunk_id: int,
) -> List[Chunk]:
    out: List[Chunk] = []
    chunk_id = start_chunk_id
    seen_rows: set[str] = set()
    per_section_count: dict[str, int] = {}
    for idx, b in enumerate(blocks):
        if b.block_type != "table_row":
            continue
        if not re.search(r"\d", b.text):
            continue
        if b.text in seen_rows:
            continue
        section_key = b.section_path or "-"
        if per_section_count.get(section_key, 0) >= 30:
            continue

        prev_ctx = ""
        next_ctx = ""
        for j in range(idx - 1, max(-1, idx - 4), -1):
            if blocks[j].block_type != "table_row" and blocks[j].section_path == b.section_path:
                prev_ctx = blocks[j].text
                break
        for j in range(idx + 1, min(len(blocks), idx + 4)):
            if blocks[j].block_type != "table_row" and blocks[j].section_path == b.section_path:
                next_ctx = blocks[j].text
                break

        body = "\n".join(
            x
            for x in [
                "[TABLE_CONTEXT]",
                f"SECTION={b.section_path or '-'}",
                prev_ctx,
                b.text,
                next_ctx,
            ]
            if x
        ).strip()
        if len(body) < 60:
            continue
        out.append(
            Chunk(
                chunk_id=chunk_id,
                text=body,
                source_path=source_path,
                source_file=source_file,
                section_path=b.section_path,
                block_types=["table_context", "table_row"],
                chunk_role="table_context",
                parent_id=None,
            )
        )
        seen_rows.add(b.text)
        per_section_count[section_key] = per_section_count.get(section_key, 0) + 1
        chunk_id += 1
    return out


def chunk_blocks(blocks: Iterable[Block], chunk_size: int, overlap: int) -> List[Chunk]:
    block_list = list(blocks)
    if not block_list:
        return []

    parent_chunks: List[Chunk] = []
    current_texts: List[str] = []
    current_types: List[str] = []
    current_section = ""
    current_source_path = block_list[0].source_path
    current_source_file = block_list[0].source_file
    chunk_id = 0

    def flush_parent() -> None:
        nonlocal chunk_id, current_texts, current_types, current_section, current_source_path, current_source_file
        if not current_texts:
            return
        text = "\n".join(current_texts).strip()
        if not text:
            return
        parent_chunks.append(
            Chunk(
                chunk_id=chunk_id,
                text=text,
                source_path=current_source_path,
                source_file=current_source_file,
                section_path=current_section,
                block_types=sorted(set(current_types)),
                chunk_role="parent",
                parent_id=None,
            )
        )
        chunk_id += 1

        has_table = "table_row" in current_types
        if overlap > 0 and not has_table:
            joined = "\n".join(current_texts)
            tail = joined[-overlap:]
            current_texts = [tail] if tail else []
            current_types = ["overlap"] if tail else []
        else:
            current_texts = []
            current_types = []

    for block in block_list:
        if not current_texts:
            current_source_path = block.source_path
            current_source_file = block.source_file
            current_section = block.section_path
        elif block.source_path != current_source_path or block.section_path != current_section:
            flush_parent()
            current_source_path = block.source_path
            current_source_file = block.source_file
            current_section = block.section_path

        if current_types:
            prev_is_table = "table_row" in current_types
            curr_is_table = block.block_type == "table_row"
            if prev_is_table != curr_is_table:
                flush_parent()
                current_source_path = block.source_path
                current_source_file = block.source_file
                current_section = block.section_path

        effective_chunk_size = max(chunk_size, 1400) if block.block_type != "table_row" else max(1800, chunk_size * 3)
        candidate = "\n".join(current_texts + [block.text])
        if len(candidate) > effective_chunk_size and current_texts:
            flush_parent()
            current_source_path = block.source_path
            current_source_file = block.source_file
            current_section = block.section_path

        current_texts.append(block.text)
        current_types.append(block.block_type)

    flush_parent()

    child_chunks: List[Chunk] = []
    child_size = max(280, chunk_size // 3)
    child_overlap = max(60, overlap // 2)
    for parent in parent_chunks:
        if "table_row" in parent.block_types:
            continue
        if len(parent.text) < max(1400, int(chunk_size * 1.2)):
            continue
        pieces = _split_child_text(parent.text, child_size=child_size, child_overlap=child_overlap)
        for piece in pieces:
            child_chunks.append(
                Chunk(
                    chunk_id=chunk_id,
                    text=piece,
                    source_path=parent.source_path,
                    source_file=parent.source_file,
                    section_path=parent.section_path,
                    block_types=sorted(set(parent.block_types + ["child"])),
                    chunk_role="child",
                    parent_id=parent.chunk_id,
                )
            )
            chunk_id += 1

    table_context_chunks = _build_table_context_chunks(
        block_list,
        source_path=block_list[0].source_path,
        source_file=block_list[0].source_file,
        start_chunk_id=chunk_id,
    )
    if table_context_chunks:
        chunk_id = table_context_chunks[-1].chunk_id + 1

    all_chunks = parent_chunks + child_chunks + table_context_chunks
    role_counts: dict[str, int] = {}
    for c in all_chunks:
        role_counts[c.chunk_role] = role_counts.get(c.chunk_role, 0) + 1
    logger.info(
        "Chunking finished | chunks=%s | parent=%s | child=%s | table_ctx=%s | chunk_size=%s | overlap=%s",
        len(all_chunks),
        role_counts.get("parent", 0),
        role_counts.get("child", 0),
        role_counts.get("table_context", 0),
        chunk_size,
        overlap,
    )
    return all_chunks


def build_corpus(data_dir: Path | None = None) -> List[Chunk]:
    data_dir = data_dir or SETTINGS.data_dir
    docx_files = sorted(data_dir.glob("*.docx"))
    logger.info("Building corpus from %s | files=%s", data_dir, len(docx_files))

    all_chunks: List[Chunk] = []
    offset = 0
    for path in docx_files:
        blocks = parse_docx(path)
        chunks = chunk_blocks(blocks, SETTINGS.chunk_size, SETTINGS.chunk_overlap)
        logger.info("File processed: %s | chunks=%s", path.name, len(chunks))
        for chunk in chunks:
            chunk.chunk_id += offset
        offset += len(chunks)
        all_chunks.extend(chunks)

    logger.info("Corpus ready | total_chunks=%s", len(all_chunks))
    return all_chunks
