import os
import sys
import glob
import re
import json
from typing import List, Tuple, Dict
import gc
import numpy as np
import os
from docx import Document 
import faiss
import torch
from tqdm import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter
from datetime import datetime
from input_data import QUESTIONS, ANSWERS
from retriever import Retriever
from models import Chunk
import logging
import torch
from generator import Generator
import json


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Решение конфликта OpenMP на macOS
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

DATA_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), os.pardir, "тз для кандидата", "data")
)
CACHE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "cache"))
RESULTS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, "results"))


def flush():
    torch.mps.empty_cache()
    gc.collect()


def save_results_to_file(outputs: List[Dict], run_id: str = None):
    """Сохраняет результаты в папку results с уникальным именем"""
    
    if run_id is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        run_id = f"run_{timestamp}"
        
    # Создаем папку results если её нет
    if not os.path.exists(RESULTS_DIR):
        os.makedirs(RESULTS_DIR)
        print(f"Создана папка для результатов: {RESULTS_DIR}")
        
    # Сохраняем в папку results с уникальным именем
    results_out_path = os.path.join(RESULTS_DIR, f"{run_id}.json")
    
    with open(results_out_path, "w", encoding="utf-8") as f:
        json.dump(outputs, f, ensure_ascii=False, indent=2)
    
    logger.info(f"Результаты сохранены в {results_out_path}")
    
    return results_out_path


def read_docx_text(docx_path: str) -> str:
    """
    Читает DOCX файл и возвращает текст
    
    Returns:
        str: текст документа
    """
    document = Document(docx_path)
    parts: List[str] = []

    # Обрабатываем параграфы
    for paragraph in document.paragraphs:
        txt = paragraph.text.strip()
        if txt:
            parts.append(txt)
    
    # Обрабатываем таблицы
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    text = "\n".join(parts)
    text = re.sub(r"\n{2,}", "\n", text)
    return text


def split_text_into_chunks(text: str, max_chars: int = 900, overlap: int = 200) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_chars,
        chunk_overlap=overlap,
    )
    docs = splitter.split_text(text)
    return docs


def build_corpus(data_dir: str) -> List[Chunk]:
    
    docx_files = sorted(glob.glob(os.path.join(data_dir, "*.docx")))
    corpus: List[Chunk] = []
    chunk_id = 0
    
    for path in tqdm(docx_files, desc="Чтение DOCX"):
        try:
            text = read_docx_text(path)
            
            # Разбиваем на чанки
            chunks = split_text_into_chunks(text)
            
            # Создаем чанки без изображений
            for chunk_text in chunks:
                corpus.append(Chunk(
                    text=chunk_text, 
                    source_path=path, 
                    chunk_id=chunk_id
                ))
                chunk_id += 1
                            
        except Exception as e:
            logger.error(f"Ошибка при обработке файла {path}: {e}")
            continue
    
    return corpus


def answer_questions(question=None):
    
    logger.info(f"Папка с данными: {DATA_DIR}")
    
    corpus = build_corpus(DATA_DIR)
    if not corpus:
        logger.warning("Не удалось собрать корпус из DOCX.")
        return
    
    logger.info(f"Корпус построен: {len(corpus)} чанков")

    retriever = Retriever() 
    generator = Generator()

    # Проверяем, нужно ли принудительно пересчитать индекс
    force_rebuild = "--rebuild" in sys.argv
    retriever.build(corpus, use_cache=not force_rebuild)

    outputs = []
    if not question:
        for i, q in enumerate(QUESTIONS, start=1):
            logger.info(f'\nОбработка вопроса {i}')
            
            # Поиск релевантных документов
            retrieved = retriever.search(q)
            # Генерация ответа
            ans, srcs = generator.generate_answer(q, retrieved)
            
            outputs.append({
                "question": q, 
                "answer": ans, 
                "sources": srcs,
            })

        # Валидация всех ответов
        logger.info("\n\nПроводим валидацию ответов...")
        generated_answers = [item["answer"] for item in outputs]
        validation_results = generator.validate_answers_batch(QUESTIONS, generated_answers, ANSWERS)
        
        # Добавляем результаты валидации к каждому ответу
        for i, item in enumerate(outputs):
            if i < len(validation_results):
                validation = validation_results[i]
                item["validation"] = {
                    "score": validation.score,
                    "reference_answer": ANSWERS[i]
                }
            else:
                item["validation"] = {
                    "score": 0.0,
                    "reference_answer": ANSWERS[i]
                }
    else:
        retrieved = retriever.search(question)
        ans, srcs = generator.generate_answer(question, retrieved)
        outputs.append({
            "question": question, 
            "answer": ans, 
            "sources": srcs,
        })
        
    save_results_to_file(outputs)


def main():
    force_rebuild = False
    
    if force_rebuild:
        logger.info("Принудительный пересчет индекса...")
        # Удаляем кэш
        import shutil
        if os.path.exists(CACHE_DIR):
            shutil.rmtree(CACHE_DIR)
            logger.info("Кэш удален")

    # question = "В каких зонах по весу снежного покрова находятся Мурманск?"
    # answer_questions(question)
    
    # answer_questions()
    res = json.load(open(os.path.join(RESULTS_DIR, "run_20250817_184742.json")))
    for item in res:
        
        print('Вопрос: ',item["question"])
        print('Ответ: ',item["answer"])
        print("-"*100)


if __name__ == "__main__":
    main()


