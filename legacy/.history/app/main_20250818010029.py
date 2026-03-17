#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG система для анализа технических документов
==============================================

Основной модуль для обработки DOCX документов и ответов на вопросы
с использованием гибридного поиска (FAISS + BM25) и LLM генерации.

Архитектура:
    1. Предобработка данных (извлечение текста из DOCX)
    2. Чанкинг документов на фрагменты
    3. Построение индексов (FAISS для семантики + BM25 для лексики)
    4. Гибридный поиск релевантных фрагментов
    5. Генерация ответов с помощью LLM
    6. Валидация качества ответов
    7. Сохранение результатов

Автор: AI Assistant
Дата: 2024
"""

import os
import sys
import glob
import re
import json
import shutil
from typing import List, Tuple, Dict, Optional
import gc
import logging
from datetime import datetime

import numpy as np
import torch
import faiss
from docx import Document 
from tqdm import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Локальные импорты
from input_data import QUESTIONS, ANSWERS
from retriever import Retriever
from models import Chunk
from generator import Generator

# =============================================================================
# КОНФИГУРАЦИЯ И НАСТРОЙКИ
# =============================================================================

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Решение конфликта OpenMP на macOS
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

# Пути к данным и кэшу
DATA_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), os.pardir, "тз для кандидата", "data")
)
CACHE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "cache"))
RESULTS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, "results"))

# Параметры чанкинга
DEFAULT_CHUNK_SIZE = 900  # Оптимальный размер чанка в символах
DEFAULT_CHUNK_OVERLAP = 200  # Перекрытие между чанками для связности

# =============================================================================
# УТИЛИТЫ ДЛЯ УПРАВЛЕНИЯ ПАМЯТЬЮ
# =============================================================================

def flush_memory() -> None:
    """
    Очищает кэш GPU и освобождает память Python.
    
    Важно для работы с большими моделями на ограниченной памяти.
    """
    if torch.backends.mps.is_available():
        torch.mps.empty_cache()
    elif torch.cuda.is_available():
        torch.cuda.empty_cache()
    gc.collect()
    logger.debug("Память очищена")

# =============================================================================
# СОХРАНЕНИЕ И ВЫВОД РЕЗУЛЬТАТОВ
# =============================================================================

def save_results_to_file(outputs: List[Dict], run_id: Optional[str] = None) -> str:
    """
    Сохраняет результаты анализа в JSON файл с уникальным именем.
    
    Args:
        outputs: Список результатов для каждого вопроса
        run_id: Уникальный идентификатор запуска (генерируется автоматически)
        
    Returns:
        str: Путь к сохранённому файлу
        
    Raises:
        OSError: При ошибках записи файла
    """
    try:
        # Генерируем уникальный ID на основе времени
        if run_id is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            run_id = f"run_{timestamp}"
            
        # Создаем папку results если её нет
        if not os.path.exists(RESULTS_DIR):
            os.makedirs(RESULTS_DIR)
            logger.info(f"Создана папка для результатов: {RESULTS_DIR}")
            
        # Формируем путь к файлу результатов
        results_path = os.path.join(RESULTS_DIR, f"{run_id}.json")
        
        # Сохраняем с правильной кодировкой для русских символов
        with open(results_path, "w", encoding="utf-8") as f:
            json.dump(outputs, f, ensure_ascii=False, indent=2)
        
        logger.info(f"✅ Результаты сохранены: {results_path}")
        return results_path
        
    except Exception as e:
        logger.error(f"❌ Ошибка сохранения результатов: {e}")
        raise


def read_docx_text(docx_path: str) -> str:
    """
    Извлекает весь текстовый контент из DOCX файла.
    
    Args:
        docx_path: Путь к DOCX файлу
        
    Returns:
        str: Объединённый текст документа
        
    Raises:
        Exception: При ошибках чтения файла
        
    Извлекает:
        - Все параграфы документа
        - Содержимое таблиц (построчно через разделитель "|")
        - Убирает избыточные переводы строк
    """
    try:
        document = Document(docx_path)
        text_parts: List[str] = []

        # Извлекаем текст из всех параграфов
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:  # Пропускаем пустые параграфы
                text_parts.append(text)
        
        # Извлекаем содержимое таблиц
        for table in document.tables:
            for row in table.rows:
                # Объединяем ячейки строки через разделитель
                row_cells = [
                    cell.text.strip() 
                    for cell in row.cells 
                    if cell.text and cell.text.strip()
                ]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        # Объединяем все части и убираем лишние переводы строк
        full_text = "\n".join(text_parts)
        clean_text = re.sub(r"\n{2,}", "\n", full_text)  # Максимум один перевод строки
        
        return clean_text
        
    except Exception as e:
        logger.error(f"❌ Ошибка чтения DOCX файла {docx_path}: {e}")
        raise


def split_text_into_chunks(
    text: str, 
    max_chars: int = DEFAULT_CHUNK_SIZE, 
    overlap: int = DEFAULT_CHUNK_OVERLAP
) -> List[str]:
    """
    Разбивает текст на семантически связанные чанки.
    
    Args:
        text: Исходный текст для разбиения
        max_chars: Максимальный размер чанка в символах
        overlap: Размер перекрытия между соседними чанками
        
    Returns:
        List[str]: Список текстовых чанков
        
    Использует RecursiveCharacterTextSplitter для сохранения
    семантической целостности при разбиении.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_chars,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ".", "!", "?", ";", ",", " ", ""]
    )
    
    chunks = splitter.split_text(text)
    logger.debug(f"Текст разбит на {len(chunks)} чанков")
    
    return chunks


def build_corpus(data_dir: str) -> List[Chunk]:
    """
    Строит корпус чанков из всех DOCX файлов в указанной папке.
    
    Args:
        data_dir: Путь к папке с DOCX документами
        
    Returns:
        List[Chunk]: Список объектов Chunk с текстом и метаданными
        
    Raises:
        ValueError: Если папка не существует или пуста
        
    Процесс:
        1. Сканирует папку на наличие DOCX файлов
        2. Извлекает текст из каждого файла
        3. Разбивает на чанки с перекрытием
        4. Создаёт объекты Chunk с уникальными ID
    """
    if not os.path.exists(data_dir):
        raise ValueError(f"Папка с данными не существует: {data_dir}")
    
    # Ищем все DOCX файлы в папке
    docx_pattern = os.path.join(data_dir, "*.docx")
    docx_files = sorted(glob.glob(docx_pattern))
    
    if not docx_files:
        raise ValueError(f"DOCX файлы не найдены в папке: {data_dir}")
    
    logger.info(f"📁 Найдено DOCX файлов: {len(docx_files)}")
    
    corpus: List[Chunk] = []
    chunk_id = 0
    
    # Обрабатываем каждый файл с прогресс-баром
    for file_path in tqdm(docx_files, desc="🔄 Обработка DOCX"):
        try:
            # Извлекаем текст
            text = read_docx_text(file_path)
            
            if not text.strip():
                logger.warning(f"⚠️ Пустой файл: {os.path.basename(file_path)}")
                continue
                
            # Разбиваем на чанки
            text_chunks = split_text_into_chunks(text)
            
            # Создаём объекты Chunk
            for chunk_text in text_chunks:
                if chunk_text.strip():  # Пропускаем пустые чанки
                    corpus.append(Chunk(
                        text=chunk_text.strip(), 
                        source_path=file_path, 
                        chunk_id=chunk_id
                    ))
                    chunk_id += 1
                    
            logger.debug(f"✅ {os.path.basename(file_path)}: {len(text_chunks)} чанков")
                            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки {file_path}: {e}")
            continue
    
    logger.info(f"📚 Корпус построен: {len(corpus)} чанков из {len(docx_files)} файлов")
    return corpus

# =============================================================================
# ОСНОВНАЯ ЛОГИКА RAG СИСТЕМЫ
# =============================================================================

def answer_questions(question: Optional[str] = None) -> List[Dict]:
    """
    Основная функция для ответов на вопросы с использованием RAG.
    
    Args:
        question: Конкретный вопрос (если None, обрабатывает все из QUESTIONS)
        
    Returns:
        List[Dict]: Результаты анализа с ответами и метаданными
        
    Pipeline:
        1. Построение корпуса из DOCX документов
        2. Создание и обучение retriever'а
        3. Поиск релевантных чанков для каждого вопроса
        4. Генерация ответов с помощью LLM
        5. Валидация качества ответов
        6. Сохранение и вывод результатов
    """
    logger.info(f"🚀 Запуск RAG системы")
    logger.info(f"📂 Папка с данными: {DATA_DIR}")
    
    # Этап 1: Построение корпуса документов
    try:
        corpus = build_corpus(DATA_DIR)
        if not corpus:
            logger.error("❌ Не удалось построить корпус документов")
            return []
            
    except Exception as e:
        logger.error(f"❌ Ошибка построения корпуса: {e}")
        return []

    # Этап 2: Инициализация компонентов
    logger.info("🔧 Инициализация retriever и generator...")
    retriever = Retriever()
    generator = Generator()

    # Этап 3: Построение индексов
    try:
        force_rebuild = "--rebuild" in sys.argv
        if force_rebuild:
            logger.info("🔄 Принудительное обновление индексов...")
            
        retriever.build(corpus, use_cache=not force_rebuild)
        
    except Exception as e:
        logger.error(f"❌ Ошибка построения индексов: {e}")
        return []

    # Этап 4: Обработка вопросов
    outputs = []
    
    if question is None:
        # Обрабатываем все вопросы из списка
        questions_to_process = QUESTIONS
        logger.info(f"📝 Обработка {len(questions_to_process)} вопросов...")
    else:
        # Обрабатываем один конкретный вопрос
        questions_to_process = [question]
        logger.info(f"📝 Обработка одного вопроса...")

    # Генерируем ответы для каждого вопроса
    for i, q in enumerate(questions_to_process, start=1):
        logger.info(f'\n🔍 Обработка вопроса {i}/{len(questions_to_process)}')
        logger.debug(f"Вопрос: {q}")
        
        try:
            # Поиск релевантных документов
            retrieved_chunks = retriever.search(q)
            logger.info(f"📚 Найдено релевантных чанков: {len(retrieved_chunks)}")
            
            # Генерация ответа
            answer, sources = generator.generate_answer(q, retrieved_chunks)
            
            # Сохраняем результат
            outputs.append({
                "question": q, 
                "answer": answer, 
                "sources": sources,
            })
            
            # Очищаем память после каждого вопроса
            flush_memory()
            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки вопроса {i}: {e}")
            # Добавляем пустой результат для сохранения порядка
            outputs.append({
                "question": q,
                "answer": f"Ошибка обработки: {str(e)}",
                "sources": [],
            })

    # Этап 5: Валидация ответов (только для полного списка вопросов)
    if question is None and outputs:
        logger.info("\n🔍 Проведение валидации ответов...")
        try:
            generated_answers = [item["answer"] for item in outputs]
            validation_results = generator.validate_answers_batch(
                QUESTIONS, generated_answers, ANSWERS
            )
            
            # Добавляем результаты валидации
            for i, output_item in enumerate(outputs):
                if i < len(validation_results):
                    validation = validation_results[i]
                    output_item["validation"] = {
                        "score": validation.score,
                        "reference_answer": ANSWERS[i]
                    }
                else:
                    # Если валидация не удалась
                    output_item["validation"] = {
                        "score": 0.0,
                        "reference_answer": ANSWERS[i] if i < len(ANSWERS) else "Нет эталона"
                    }
                    
        except Exception as e:
            logger.error(f"❌ Ошибка валидации: {e}")
    
    # Этап 6: Сохранение и вывод результатов
    if outputs:
        try:
            results_path = save_results_to_file(outputs)
            print_results_summary(outputs)
            
        except Exception as e:
            logger.error(f"❌ Ошибка сохранения результатов: {e}")
    
    logger.info("🎉 Обработка завершена!")
    return outputs

# =============================================================================
# УПРАВЛЕНИЕ КЭШЕМ
# =============================================================================

def clear_cache() -> None:
    """
    Очищает кэш индексов для принудительного пересчёта.
    
    Удаляет всю папку cache/, что приведёт к пересозданию
    FAISS индексов при следующем запуске.
    """
    try:
        if os.path.exists(CACHE_DIR):
            shutil.rmtree(CACHE_DIR)
            logger.info(f"🗑️ Кэш очищен: {CACHE_DIR}")
        else:
            logger.info("ℹ️ Кэш уже пуст")
            
    except Exception as e:
        logger.error(f"❌ Ошибка очистки кэша: {e}")

# =============================================================================
# ТОЧКА ВХОДА
# =============================================================================

def main() -> None:
    """
    Главная функция программы.
    
    Поддерживаемые режимы работы:
        - python main.py                    # Обработка всех вопросов
        - python main.py --rebuild          # С принудительным обновлением индексов
        - python main.py --clear-cache      # Только очистка кэша
        - python main.py --load-results     # Загрузка последних результатов
    """
    logger.info("🎯 Запуск RAG системы для анализа технических документов")
    
    # Обработка аргументов командной строки
    if "--clear-cache" in sys.argv:
        clear_cache()
        return
        
    if "--load-results" in sys.argv:
        # Загружаем последний файл результатов
        try:
            result_files = [f for f in os.listdir(RESULTS_DIR) if f.startswith("run_")]
            if result_files:
                latest_file = sorted(result_files)[-1]
                load_and_print_results(latest_file)
            else:
                logger.warning("⚠️ Файлы результатов не найдены")
        except Exception as e:
            logger.error(f"❌ Ошибка загрузки результатов: {e}")
        return
    
    # Принудительная очистка кэша при --rebuild
    if "--rebuild" in sys.argv:
        logger.info("🔄 Режим принудительного обновления...")
        clear_cache()
    
    # Основная обработка
    try:
        # Можно раскомментировать для тестирования одного вопроса:
        # test_question = "В каких зонах по весу снежного покрова находятся Херсон и Мелитополь?"
        # answer_questions(test_question)
        
        # Обработка всех вопросов
        answer_questions()
        
    except KeyboardInterrupt:
        logger.info("⚠️ Обработка прервана пользователем")
    except Exception as e:
        logger.error(f"❌ Критическая ошибка: {e}")
        raise


if __name__ == "__main__":
    main()